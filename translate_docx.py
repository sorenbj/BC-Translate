#!/usr/bin/env python3
"""
BC-Translate: Word Document Translation Tool
Translates Word documents from English to Danish while preserving formatting.
"""

import argparse
import os
import sys
from docx import Document
from deep_translator import GoogleTranslator


def translate_text(text, source_lang='en', target_lang='da'):
    """
    Translate text from source language to target language.
    
    Args:
        text: Text to translate
        source_lang: Source language code (default: 'en' for English)
        target_lang: Target language code (default: 'da' for Danish)
    
    Returns:
        Translated text
    """
    if not text or not text.strip():
        return text
    
    try:
        translator = GoogleTranslator(source=source_lang, target=target_lang)
        # Handle long text by splitting if necessary (Google Translate has character limits)
        max_length = 4500
        if len(text) > max_length:
            # Split by sentences/paragraphs and translate in chunks
            chunks = []
            current_chunk = ""
            for sentence in text.split('. '):
                if len(current_chunk) + len(sentence) < max_length:
                    current_chunk += sentence + '. '
                else:
                    if current_chunk:
                        chunks.append(current_chunk)
                    current_chunk = sentence + '. '
            if current_chunk:
                chunks.append(current_chunk)
            
            translated_chunks = [translator.translate(chunk) for chunk in chunks]
            return ''.join(translated_chunks)
        else:
            return translator.translate(text)
    except Exception as e:
        print(f"Warning: Failed to translate text: {e}")
        return text


def translate_paragraph(paragraph, source_lang='en', target_lang='da'):
    """
    Translate a paragraph while preserving formatting (runs).
    
    Args:
        paragraph: python-docx paragraph object
        source_lang: Source language code
        target_lang: Target language code
    """
    if not paragraph.text.strip():
        return
    
    # Collect all text from runs
    full_text = paragraph.text
    
    # Translate the full text
    translated_text = translate_text(full_text, source_lang, target_lang)
    
    # If paragraph has runs with formatting, we need to be careful
    if len(paragraph.runs) > 0:
        # Simple approach: replace text in first run, clear others
        # This preserves some formatting but may lose complex formatting
        # For better preservation, we'd need more sophisticated run handling
        paragraph.runs[0].text = translated_text
        for i in range(1, len(paragraph.runs)):
            paragraph.runs[i].text = ''
    else:
        paragraph.text = translated_text


def translate_table(table, source_lang='en', target_lang='da'):
    """
    Translate all text in a table while preserving structure.
    
    Args:
        table: python-docx table object
        source_lang: Source language code
        target_lang: Target language code
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                translate_paragraph(paragraph, source_lang, target_lang)


def translate_document(input_path, output_path, source_lang='en', target_lang='da'):
    """
    Translate a Word document while preserving formatting.
    
    Args:
        input_path: Path to input .docx file
        output_path: Path to output .docx file
        source_lang: Source language code (default: 'en')
        target_lang: Target language code (default: 'da')
    """
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"Input file not found: {input_path}")
    
    print(f"Loading document: {input_path}")
    doc = Document(input_path)
    
    print(f"Translating from {source_lang} to {target_lang}...")
    
    # Translate paragraphs
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            print(f"  Translating paragraph {i+1}/{len(doc.paragraphs)}")
            translate_paragraph(paragraph, source_lang, target_lang)
    
    # Translate tables
    for i, table in enumerate(doc.tables):
        print(f"  Translating table {i+1}/{len(doc.tables)}")
        translate_table(table, source_lang, target_lang)
    
    # Save translated document
    print(f"Saving translated document: {output_path}")
    doc.save(output_path)
    print("Translation complete!")


def main():
    """Main entry point for the CLI."""
    parser = argparse.ArgumentParser(
        description='Translate Word documents from English to Danish while preserving formatting.',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
Examples:
  %(prog)s input.docx output.docx
  %(prog)s input.docx output.docx --source en --target da
        '''
    )
    
    parser.add_argument('input', help='Input .docx file path')
    parser.add_argument('output', help='Output .docx file path', nargs='?')
    parser.add_argument('--source', '-s', default='en', 
                        help='Source language code (default: en)')
    parser.add_argument('--target', '-t', default='da',
                        help='Target language code (default: da)')
    
    args = parser.parse_args()
    
    # If output not specified, generate it from input
    if args.output is None:
        base, ext = os.path.splitext(args.input)
        args.output = f"{base}_translated{ext}"
    
    try:
        translate_document(args.input, args.output, args.source, args.target)
        return 0
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1


if __name__ == '__main__':
    sys.exit(main())
